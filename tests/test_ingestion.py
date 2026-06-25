"""Tests for multi-format ingestion (#8). Network/binary-dependent formats
(URL fetch, YouTube, OCR) are tested at the parsing layer, not over the wire."""

import io

import pytest

from src.ingestion import (
    SUPPORTED_UPLOAD_EXTENSIONS,
    _youtube_id,
    extract_pages,
    html_to_text,
)


def test_txt_and_md_ingestion_preserves_lines():
    pages = extract_pages("notes.txt", b"Heading One\nsome body text\n\nHeading Two\nmore text")
    assert len(pages) == 1
    assert pages[0]["source"] == "notes.txt"
    assert "Heading One" in pages[0]["text"]
    assert "\n" in pages[0]["text"]  # line structure kept for heading-aware chunking

    md = extract_pages("readme.md", b"# Title\nbody")
    assert md[0]["text"].startswith("# Title")


def test_docx_ingestion():
    import docx

    document = docx.Document()
    document.add_paragraph("First paragraph about retrieval.")
    document.add_paragraph("Second paragraph about embeddings.")
    buffer = io.BytesIO()
    document.save(buffer)

    pages = extract_pages("report.docx", buffer.getvalue())
    assert pages
    assert "retrieval" in pages[0]["text"]
    assert "embeddings" in pages[0]["text"]


def test_unsupported_extension_raises():
    with pytest.raises(ValueError):
        extract_pages("archive.zip", b"\x00\x01")


def test_supported_extensions_include_key_formats():
    for ext in (".pdf", ".docx", ".txt", ".md", ".png", ".jpg"):
        assert ext in SUPPORTED_UPLOAD_EXTENSIONS


def test_html_to_text_strips_tags_and_scripts():
    html = "<html><head><style>.x{}</style></head><body><h1>Title</h1>" \
           "<script>ignore()</script><p>Hello world</p></body></html>"
    text = html_to_text(html)
    assert "Title" in text
    assert "Hello world" in text
    assert "ignore()" not in text
    assert ".x{}" not in text


def test_youtube_id_parsing():
    assert _youtube_id("https://www.youtube.com/watch?v=dQw4w9WgXcQ") == "dQw4w9WgXcQ"
    assert _youtube_id("https://youtu.be/dQw4w9WgXcQ") == "dQw4w9WgXcQ"
    assert _youtube_id("dQw4w9WgXcQ") == "dQw4w9WgXcQ"
    with pytest.raises(ValueError):
        _youtube_id("https://example.com/no-video")
